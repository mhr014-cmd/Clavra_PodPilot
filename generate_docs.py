"""Generate all Clavra ProdPilot docs: summary.docx, structure.docx, structure.md, SRS.pdf"""
import os
from pathlib import Path

OUT = Path("D:/projects/Claude/Clavra_ProdPilot/docs")
OUT.mkdir(exist_ok=True)

# ═══════════════════════════════════════════════════════════════
# DOCX helpers
# ═══════════════════════════════════════════════════════════════
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def shade_cell(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shd)

def hdr_row(table, headers, fill="1E40AF"):
    row = table.rows[0]
    for i, h in enumerate(headers):
        row.cells[i].text = h
        for run in row.cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
        shade_cell(row.cells[i], fill)

def add_row(table, *vals, key_bold=True, key_shade="DCE8F8"):
    row = table.add_row()
    for i, v in enumerate(vals):
        row.cells[i].text = str(v)
    if key_bold:
        for run in row.cells[0].paragraphs[0].runs:
            run.bold = True
    if key_shade:
        shade_cell(row.cells[0], key_shade)
    return row

def dset_heading(doc, text, level=1, rgb=(30, 64, 175)):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(*rgb)
    return h


# ═══════════════════════════════════════════════════════════════
# 1. PROJECT SUMMARY — DOCX
# ═══════════════════════════════════════════════════════════════
doc = Document()

title = doc.add_heading("Clavra ProdPilot™", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in title.runs:
    r.font.color.rgb = RGBColor(30, 64, 175)
    r.font.size = Pt(28)

sub = doc.add_paragraph("AI-Powered Manufacturing ERP — Project Summary  |  v1.0  |  June 2026")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in sub.runs:
    r.font.color.rgb = RGBColor(100, 116, 139)
    r.italic = True
doc.add_paragraph()

dset_heading(doc, "1. Brief Summary")
doc.add_paragraph(
    "Clavra ProdPilot™ is a full-stack, AI-driven ERP system purpose-built for garment manufacturing. "
    "It replaces WhatsApp and Excel-based workflows with a real-time dashboard, an intelligent AI Copilot "
    "(GPT-4o + Ollama), and a RAG-powered Knowledge Base — enabling factory managers to query production "
    "status, monitor shipments, manage inventory, and retrieve SOPs using plain natural language or voice."
)

dset_heading(doc, "2. Problem Statement")
doc.add_paragraph(
    "Small and mid-size garment factories in South & South-East Asia rely on WhatsApp messages and "
    "Excel sheets to manage production orders, shipments, and raw material inventory. This leads to "
    "delayed decisions, missed deadlines, stock-outs, and zero traceability. Clavra ProdPilot™ delivers "
    "an affordable, AI-first ERP that requires no training — workers simply ask questions in plain language."
)

dset_heading(doc, "3. Key Features")
features = [
    ("AI Copilot (18+ Intents)", "GPT-4o intent classification with keyword + Ollama fallbacks; voice input via Web Speech API"),
    ("Production Management", "Track orders, progress (produced vs target), defect qty, delivery dates with inline editing"),
    ("Shipment Tracking", "Real-time status lifecycle linked to production orders"),
    ("RAG Knowledge Base", "Upload PDFs/DOCX/XLSX SOPs; semantic vector search answers with citations"),
    ("Inventory Management", "Material stock, reorder alerts, supplier info"),
    ("Production Lines", "Line A/B/C/D assignment, status, and efficiency tracking"),
    ("Quality Control", "Defect analysis, QC records, inspection reports"),
    ("Analytics Dashboard", "KPI cards, order trend charts, inventory analytics"),
    ("Multi-tenant RBAC", "Org-scoped data with Admin/Manager/Operator roles"),
    ("WebSocket Streaming", "Real-time AI responses streamed token-by-token with intent badges"),
]
t = doc.add_table(rows=1, cols=2)
t.style = "Table Grid"
hdr_row(t, ["Feature", "Description"])
for name, desc in features:
    add_row(t, name, desc)

doc.add_paragraph()
dset_heading(doc, "4. Technology Stack")
stack = [
    ("Backend", "Python 3.11+, FastAPI, SQLAlchemy 2.0 (async), Alembic, Uvicorn"),
    ("Database", "PostgreSQL 16 + pgvector (1536-dim vector embeddings)"),
    ("Frontend", "React 18, TypeScript, Vite 5, Tailwind CSS v4, Zustand"),
    ("AI / ML", "OpenAI GPT-4o (intent + RAG), Ollama llama3.1:8b (local fallback), nomic-embed-text"),
    ("Auth", "JWT (python-jose), bcrypt, refresh token rotation"),
    ("Voice", "Web Speech API (browser-native STT + TTS)"),
    ("DevOps", "Docker Compose, Nginx reverse proxy, Alembic migrations"),
]
t2 = doc.add_table(rows=1, cols=2)
t2.style = "Table Grid"
hdr_row(t2, ["Layer", "Technologies"], fill="0F766E")
for layer, tech in stack:
    add_row(t2, layer, tech, key_shade="CCFBF1")

doc.add_paragraph()
dset_heading(doc, "5. Architecture Overview")
doc.add_paragraph(
    "Three-tier architecture:\n"
    "  Browser (React/TS) ←→ FastAPI REST + WebSocket ←→ PostgreSQL + pgvector\n\n"
    "AI intent pipeline (three stages):\n"
    "  Stage 1 — detect_intent(): GPT-4o (8s timeout) → keyword rules → Ollama fallback\n"
    "  Stage 2 — route_intent(): 4 override guards correct misclassifications → dispatch\n"
    "  Stage 3 — Handler: service layer calls DB / RAG / SQL generator\n\n"
    "WebSocket sends intent badge immediately after Stage 1, then streams handler tokens."
)

dset_heading(doc, "6. Development Access")
for k, v in [("Frontend", "http://localhost:5174"), ("Backend API", "http://localhost:8000"),
             ("API Docs", "http://localhost:8000/docs"), ("Admin Login", "admin@clavra.com / Admin@123"),
             ("Ollama", "http://localhost:11434"), ("Org ID", "2")]:
    p = doc.add_paragraph()
    p.add_run(f"{k}: ").bold = True
    p.add_run(v)

doc.save(OUT / "project_summary.docx")
print("[OK]  project_summary.docx")


# ═══════════════════════════════════════════════════════════════
# 2. PROJECT STRUCTURE — MD
# ═══════════════════════════════════════════════════════════════
structure_md = """\
# Clavra ProdPilot™ — Project Structure

```
clavra/
├── backend/                          # FastAPI Python backend
│   ├── app/
│   │   ├── ai/                       # AI pipeline
│   │   │   ├── intent_engine.py      # Intent detection (GPT-4o → keywords → Ollama)
│   │   │   ├── intent_router.py      # Intent routing + 4 pre-routing override guards
│   │   │   ├── rag_service.py        # RAG pipeline (embed → vector search → answer)
│   │   │   ├── ollama_service.py     # Local Ollama LLM (async thread pool wrapper)
│   │   │   ├── sql_generator.py      # Text-to-SQL for analytics queries
│   │   │   ├── vision_service.py     # Image defect analysis (moondream)
│   │   │   ├── voice_service.py      # Voice transcription helpers
│   │   │   ├── tool_calling.py       # OpenAI function calling bridge
│   │   │   └── provider.py           # AI provider abstraction layer
│   │   ├── core/
│   │   │   ├── security.py           # JWT creation/validation, bcrypt hashing
│   │   │   ├── permissions.py        # RBAC permission checks
│   │   │   └── constants.py          # App-wide enums and constants
│   │   ├── models/                   # SQLAlchemy ORM models
│   │   │   ├── user.py               # User + Organisation
│   │   │   ├── production.py         # ProductionOrder
│   │   │   ├── production_line.py    # ProductionLine
│   │   │   ├── shipment.py           # Shipment
│   │   │   ├── inventory.py          # InventoryItem
│   │   │   ├── quality.py            # QCRecord
│   │   │   ├── knowledge_document.py # Uploaded SOP/policy documents
│   │   │   ├── document_chunk.py     # pgvector-embedded document chunks
│   │   │   ├── ai_chat.py            # Chat session + message history
│   │   │   ├── intent_audit.py       # Intent classification audit log
│   │   │   ├── notification.py       # In-app notifications
│   │   │   ├── planning.py           # Production planning records
│   │   │   └── vision_analysis.py    # Vision/defect analysis results
│   │   ├── routes/                   # FastAPI routers (one file per domain)
│   │   │   ├── ai_chat.py            # WebSocket AI Copilot endpoint
│   │   │   ├── ai_routes.py          # REST AI endpoints (image upload etc.)
│   │   │   ├── auth_routes.py        # /auth/login, /auth/refresh, /auth/me
│   │   │   ├── orders.py             # CRUD + link-shipment for production orders
│   │   │   ├── shipment.py           # Shipment CRUD
│   │   │   ├── inventory.py          # Inventory CRUD
│   │   │   ├── production_lines.py   # Production line management
│   │   │   ├── production.py         # Production summary endpoints
│   │   │   ├── analytics.py          # Order/production analytics
│   │   │   ├── inventory_analytics.py# Inventory trend analytics
│   │   │   ├── knowledge_routes.py   # Document upload + RAG query
│   │   │   ├── quality_routes.py     # QC records
│   │   │   └── notification.py       # Notifications
│   │   ├── schemas/                  # Pydantic v2 request/response models
│   │   ├── services/                 # Business logic layer
│   │   │   ├── ai_service.py         # AI orchestration
│   │   │   ├── auth_service.py       # Auth business logic
│   │   │   ├── production_service.py # Production order logic
│   │   │   ├── shipment_service.py   # Shipment logic
│   │   │   ├── inventory_service.py  # Inventory logic
│   │   │   ├── quality_service.py    # Quality control logic
│   │   │   └── summary_service.py    # Factory summary aggregation
│   │   ├── utils/
│   │   │   └── logger.py             # Structured JSON logging
│   │   ├── config.py                 # Pydantic settings (reads .env)
│   │   ├── database.py               # Async SQLAlchemy engine + session factory
│   │   ├── dependencies.py           # FastAPI dependency injectors (get_db, get_user)
│   │   ├── init_db.py                # DB bootstrap + seed data
│   │   └── main.py                   # FastAPI app + router registration + CORS
│   ├── alembic/                      # Database migrations
│   │   └── versions/                 # 6 migration scripts
│   ├── .env                          # Environment variables (git-ignored)
│   ├── .env.example                  # Template for new developers
│   ├── requirements.txt              # Python dependencies
│   └── Dockerfile                    # Backend Docker image
│
├── frontend/                         # React 18 + TypeScript SPA
│   ├── src/
│   │   ├── pages/
│   │   │   ├── DashboardPage.tsx     # KPI dashboard with charts
│   │   │   ├── AICopilotPage.tsx     # AI chat interface (WebSocket streaming)
│   │   │   ├── ProductionPage.tsx    # Orders + inline progress/date/shipment editing
│   │   │   ├── ProductionLinePage.tsx# Line A/B/C/D status
│   │   │   ├── ShipmentPage.tsx      # Shipment tracking
│   │   │   ├── InventoryPage.tsx     # Inventory management
│   │   │   ├── QualityPage.tsx       # QC records
│   │   │   ├── KnowledgePage.tsx     # RAG document upload + query
│   │   │   └── LoginPage.tsx         # JWT auth
│   │   ├── components/
│   │   │   ├── chat/
│   │   │   │   ├── ChatBubble.tsx    # Message bubbles with markdown rendering
│   │   │   │   ├── IntentBadge.tsx   # Intent + confidence badge
│   │   │   │   ├── VoiceButton.tsx   # Web Speech API voice input button
│   │   │   │   ├── ImageUploadZone.tsx# Drag-drop for defect image analysis
│   │   │   │   ├── SourceCitation.tsx# RAG document citations
│   │   │   │   ├── SqlViewPanel.tsx  # Expandable SQL view for analytics
│   │   │   │   └── TypingIndicator.tsx# Animated "thinking" dots
│   │   │   └── ui/                   # Button, Card, Input, Badge
│   │   ├── hooks/
│   │   │   ├── useAIChat.ts          # WebSocket chat hook
│   │   │   ├── useAuth.ts            # Auth state hook
│   │   │   └── useVoice.ts           # Voice input hook
│   │   ├── store/
│   │   │   ├── authStore.ts          # Zustand auth store
│   │   │   └── chatStore.ts          # Zustand chat store
│   │   ├── api/                      # Axios instance + interceptors
│   │   ├── services/                 # REST API + auth services
│   │   └── widgets/                  # Dashboard mini-widgets
│   ├── tailwind.config.js
│   ├── vite.config.ts
│   ├── package.json
│   └── Dockerfile
│
├── docker-compose.yml                # One-command full-stack deployment
├── nginx.conf                        # Production reverse proxy config
└── uploads/docs/                     # Knowledge base document storage
```

## Key Architecture Decisions

| Decision | Choice | Reason |
|---|---|---|
| Async framework | FastAPI + asyncpg | Non-blocking I/O for concurrent WebSocket streams |
| ORM | SQLAlchemy 2.0 async | Type-safe queries + Alembic migration support |
| Vector DB | pgvector (PostgreSQL extension) | Single DB for relational + vector — no extra service |
| State management | Zustand | Minimal boilerplate, simpler than Redux |
| Styling | Tailwind CSS v4 | Utility-first, consistent dark theme |
| AI fallback chain | GPT-4o → keywords → Ollama | Graceful degradation without OpenAI API key |
| Auth | JWT + refresh token rotation | Stateless, horizontally scalable |
| Embeddings | text-embedding-3-small (1536d) | Cost-effective; keyword fallback works without OpenAI |

## Database Schema (Key Tables)

| Table | Key Columns |
|---|---|
| users | id, email, hashed_password, role, org_id |
| organisations | id, name, plan |
| production_orders | id, order_no, buyer, status, total_qty, produced_qty, defect_qty, delivery_date, org_id |
| production_lines | id, line_name, status, current_order_id, efficiency_pct, org_id |
| shipments | id, shipment_no, status, order_id, destination, estimated_delivery, org_id |
| inventory_items | id, material_name, quantity, unit, reorder_level, supplier, org_id |
| knowledge_documents | id, original_name, doc_type, file_path, chunk_count, org_id |
| document_chunks | id, document_id, content, embedding (vector 1536), page_number |
| ai_messages | id, session_id, role, content, intent, confidence, action_type |
"""

(OUT / "project_structure.md").write_text(structure_md, encoding="utf-8")
print("[OK]  project_structure.md")


# ═══════════════════════════════════════════════════════════════
# 2b. PROJECT STRUCTURE — DOCX
# ═══════════════════════════════════════════════════════════════
doc2 = Document()
t2h = doc2.add_heading("Clavra ProdPilot™ — Project Structure", 0)
t2h.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in t2h.runs:
    r.font.color.rgb = RGBColor(30, 64, 175)
    r.font.size = Pt(22)

doc2.add_paragraph()
dset_heading(doc2, "Directory Tree", 1, rgb=(15, 118, 110))
p = doc2.add_paragraph()
rn = p.add_run(structure_md.split("```")[1].strip())
rn.font.name = "Courier New"
rn.font.size = Pt(8)

doc2.add_paragraph()
dset_heading(doc2, "Key Architecture Decisions", 1, rgb=(15, 118, 110))
arch = [
    ("Async Framework", "FastAPI + asyncpg", "Non-blocking I/O handles concurrent WebSocket streams"),
    ("ORM", "SQLAlchemy 2.0 async", "Type-safe queries with Alembic migration support"),
    ("Vector DB", "pgvector (PostgreSQL)", "Single DB for relational + vector data — no extra service needed"),
    ("Frontend State", "Zustand", "Minimal boilerplate vs Redux; simpler for auth + chat state"),
    ("Styling", "Tailwind CSS v4", "Utility-first; consistent dark factory theme"),
    ("AI Fallback Chain", "GPT-4o → keywords → Ollama", "Graceful degradation — works without OpenAI API key"),
    ("Authentication", "JWT + refresh tokens", "Stateless, horizontally scalable"),
    ("Embeddings", "text-embedding-3-small (1536d)", "Falls back to keyword search without OpenAI"),
]
t3 = doc2.add_table(rows=1, cols=3)
t3.style = "Table Grid"
hdr_row(t3, ["Decision", "Choice", "Reason"], fill="0F766E")
for a, b, c in arch:
    add_row(t3, a, b, c, key_shade="CCFBF1")

doc2.add_paragraph()
dset_heading(doc2, "Database Schema", 1, rgb=(15, 118, 110))
tables_info = [
    ("users", "id, email, hashed_password, role, org_id, last_login_at"),
    ("organisations", "id, name, plan, created_at"),
    ("production_orders", "id, order_no, buyer, status, total_qty, produced_qty, defect_qty, delivery_date, org_id"),
    ("production_lines", "id, line_name, status, current_order_id, efficiency_pct, org_id"),
    ("shipments", "id, shipment_no, status, order_id, destination, estimated_delivery, org_id"),
    ("inventory_items", "id, material_name, quantity, unit, reorder_level, supplier, org_id"),
    ("knowledge_documents", "id, original_name, doc_type, file_path, chunk_count, org_id"),
    ("document_chunks", "id, document_id, content, embedding (vector 1536), page_number"),
    ("ai_messages", "id, session_id, role, content, intent, confidence, action_type"),
]
t4 = doc2.add_table(rows=1, cols=2)
t4.style = "Table Grid"
hdr_row(t4, ["Table", "Key Columns"])
for tname, cols in tables_info:
    row = t4.add_row()
    row.cells[0].text = tname
    row.cells[1].text = cols
    for r in row.cells[0].paragraphs[0].runs:
        r.bold = True
        r.font.name = "Courier New"

doc2.save(OUT / "project_structure.docx")
print("[OK]  project_structure.docx")


# ═══════════════════════════════════════════════════════════════
# 3. SRS — PDF via reportlab
# ═══════════════════════════════════════════════════════════════
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                TableStyle, HRFlowable, PageBreak)
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY

BLUE  = colors.HexColor("#1E40AF")
TEAL  = colors.HexColor("#0F766E")
GRAY  = colors.HexColor("#64748B")
LGRAY = colors.HexColor("#F1F5F9")
DGRAY = colors.HexColor("#334155")
WHITE = colors.white

styles = getSampleStyleSheet()

S = {
    "H0":   ParagraphStyle("H0",   parent=styles["Title"],   fontSize=26, textColor=BLUE, spaceAfter=6, alignment=TA_CENTER),
    "H0s":  ParagraphStyle("H0s",  parent=styles["Normal"],  fontSize=11, textColor=GRAY, spaceAfter=20, alignment=TA_CENTER),
    "H1":   ParagraphStyle("H1",   parent=styles["Heading1"],fontSize=15, textColor=BLUE, spaceBefore=18, spaceAfter=8),
    "H2":   ParagraphStyle("H2",   parent=styles["Heading2"],fontSize=12, textColor=TEAL, spaceBefore=12, spaceAfter=6),
    "body": ParagraphStyle("body", parent=styles["Normal"],  fontSize=10, textColor=DGRAY, leading=15, spaceAfter=6, alignment=TA_JUSTIFY),
    "bul":  ParagraphStyle("bul",  parent=styles["Normal"],  fontSize=10, textColor=DGRAY, leading=14, spaceAfter=3, leftIndent=18),
    "code": ParagraphStyle("code", parent=styles["Normal"],  fontSize=8.5, fontName="Courier",
                           backColor=colors.HexColor("#F8FAFC"), leftIndent=14, rightIndent=14,
                           spaceAfter=6, leading=12, textColor=DGRAY),
}

def HR(): return HRFlowable(width="100%", thickness=1, color=colors.HexColor("#CBD5E1"), spaceAfter=6)
def B(t): return f"<b>{t}</b>"

def make_table(data, col_widths, hdr_fill=BLUE):
    t = Table(data, colWidths=col_widths)
    t.setStyle(TableStyle([
        ("BACKGROUND",  (0,0), (-1,0),  hdr_fill),
        ("TEXTCOLOR",   (0,0), (-1,0),  WHITE),
        ("FONTNAME",    (0,0), (-1,0),  "Helvetica-Bold"),
        ("ROWBACKGROUNDS",(0,1),(-1,-1),[WHITE, LGRAY]),
        ("FONTSIZE",    (0,0), (-1,-1), 9.5),
        ("GRID",        (0,0), (-1,-1), 0.5, colors.HexColor("#CBD5E1")),
        ("VALIGN",      (0,0), (-1,-1), "TOP"),
        ("PADDING",     (0,0), (-1,-1), 5),
    ]))
    return t

story = []

# ── Cover ─────────────────────────────────────────────────────
story += [
    Spacer(1, 1.5*cm),
    Paragraph("Clavra ProdPilot™", S["H0"]),
    Paragraph("Software Requirements Specification (SRS)", S["H0s"]),
    Paragraph("Version 1.0  |  June 2026  |  Confidential", S["H0s"]),
    HR(), Spacer(1, 0.5*cm),
]
doc_info = [
    [B("Document Title"), "Clavra ProdPilot™ — SRS"],
    [B("Version"), "1.0"],
    [B("Date"), "June 12, 2026"],
    [B("Status"), "Active Development"],
    [B("Stack"), "FastAPI + React 18 + PostgreSQL + pgvector + GPT-4o + Ollama"],
    [B("Target Audience"), "Developers, QA Engineers, Stakeholders"],
]
story += [
    make_table([[Paragraph(a, S["body"]), Paragraph(b, S["body"])] for a, b in doc_info],
               [4.5*cm, 13*cm], hdr_fill=colors.HexColor("#F1F5F9")),
    Spacer(1, 0.5*cm), PageBreak(),
]

# ── 1. Introduction ───────────────────────────────────────────
story += [
    Paragraph("1. Introduction", S["H1"]), HR(),
    Paragraph("1.1 Purpose", S["H2"]),
    Paragraph("This document specifies the functional and non-functional requirements for "
              "<b>Clavra ProdPilot™</b>, an AI-powered ERP for garment manufacturing. It is the "
              "authoritative reference for development, testing, and stakeholder alignment.", S["body"]),
    Paragraph("1.2 Scope", S["H2"]),
    Paragraph("The system covers: production order management, shipment tracking, raw material inventory, "
              "production line monitoring, quality control, and an AI Copilot combining structured DB "
              "lookups, SQL analytics, and RAG document retrieval from uploaded SOPs.", S["body"]),
    Paragraph("1.3 Definitions", S["H2"]),
]
defs = [
    [B("Term"), B("Definition")],
    ["ERP", "Enterprise Resource Planning — integrated management of core business processes"],
    ["RAG", "Retrieval-Augmented Generation — LLM answers grounded in retrieved documents"],
    ["Intent", "Classified purpose of a user message (e.g. shipment_status, ask_manual)"],
    ["pgvector", "PostgreSQL extension for storing and querying 1536-dim vector embeddings"],
    ["SOP", "Standard Operating Procedure — step-by-step work instruction document"],
    ["RBAC", "Role-Based Access Control — permissions tied to Admin/Manager/Operator roles"],
    ["JWT", "JSON Web Token — stateless authentication token"],
    ["Ollama", "Local LLM inference server running open-source models (llama3.1:8b)"],
]
story += [make_table([[Paragraph(a, S["body"]), Paragraph(b, S["body"])] for a, b in defs],
                     [3.5*cm, 14*cm]), Spacer(1, 0.3*cm), PageBreak()]

# ── 2. Overall Description ────────────────────────────────────
story += [
    Paragraph("2. Overall Description", S["H1"]), HR(),
    Paragraph("2.1 Product Perspective", S["H2"]),
    Paragraph("Clavra ProdPilot™ is a standalone web application replacing spreadsheet-based factory "
              "workflows. It integrates with OpenAI for cloud AI and Ollama for local AI, enabling "
              "offline-capable deployments. A REST + WebSocket API is consumed by the React SPA.", S["body"]),
    Paragraph("2.2 User Classes", S["H2"]),
]
users = [
    [B("User Class"), B("Role"), B("Primary Tasks")],
    ["Factory Admin",      "Full access", "Configure org, manage users, view all reports"],
    ["Production Manager", "Manager",     "Monitor orders, update progress, view analytics"],
    ["Line Supervisor",    "Operator",    "Update line status, record QC, log produced quantities"],
    ["Logistics Officer",  "Operator",    "Manage shipments, link orders"],
    ["AI Copilot",         "System",      "Answer natural-language queries via WebSocket"],
]
story += [make_table([[Paragraph(a, S["body"]), Paragraph(b, S["body"]), Paragraph(c, S["body"])]
                       for a, b, c in users],
                     [4.5*cm, 3.5*cm, 9.5*cm], hdr_fill=TEAL), Spacer(1, 0.3*cm)]
story += [
    Paragraph("2.3 Operating Environment", S["H2"]),
    Paragraph("• <b>Server:</b> Linux/Windows with Python 3.11+, PostgreSQL 16+, Ollama", S["bul"]),
    Paragraph("• <b>Client:</b> Chrome 120+ / Firefox 120+ / Edge 120+ with microphone for voice", S["bul"]),
    Paragraph("• <b>Docker:</b> Full-stack via docker-compose + Nginx reverse proxy", S["bul"]),
    Paragraph("• <b>Cloud AI:</b> OpenAI API (GPT-4o) — optional; graceful degradation without it", S["bul"]),
    Spacer(1, 0.3*cm), PageBreak(),
]

# ── 3. Functional Requirements ────────────────────────────────
story += [Paragraph("3. Functional Requirements", S["H1"]), HR()]

modules = [
    ("3.1 Authentication & RBAC", [
        ("FR-AUTH-01", "Login with email+password returns JWT access token + refresh token"),
        ("FR-AUTH-02", "Refresh tokens rotate on each use; expired tokens rejected with HTTP 401"),
        ("FR-AUTH-03", "All API endpoints (except /auth/login) require valid Bearer JWT"),
        ("FR-AUTH-04", "All data access scoped to org_id — no cross-org data leakage"),
        ("FR-AUTH-05", "Roles: Admin > Manager > Operator with progressively fewer write permissions"),
    ]),
    ("3.2 Production Order Management", [
        ("FR-PROD-01", "CRUD for orders: order_no, buyer, style, colour, total_qty, status, delivery_date"),
        ("FR-PROD-02", "Status lifecycle: Draft → Cutting → Sewing → Finishing → QC → Packed → Shipped"),
        ("FR-PROD-03", "Record produced_qty and defect_qty; system calculates progress percentage"),
        ("FR-PROD-04", "Inline editing of progress, delivery date, and shipment link from UI"),
        ("FR-PROD-05", "Link order to shipment via PUT /orders/{id}/link-shipment"),
    ]),
    ("3.3 Production Line Management", [
        ("FR-LINE-01", "Define lines (A/B/C/D) with capacity and current order assignment"),
        ("FR-LINE-02", "Real-time status: Idle / Running / Maintenance"),
        ("FR-LINE-03", "Track efficiency percentage per line"),
    ]),
    ("3.4 Shipment Management", [
        ("FR-SHIP-01", "CRUD for shipments: shipment_no, destination, carrier, status, estimated_delivery"),
        ("FR-SHIP-02", "Status: Pending → In Transit → Delivered → Cancelled"),
        ("FR-SHIP-03", "Link/unlink production orders to shipments"),
    ]),
    ("3.5 Inventory Management", [
        ("FR-INV-01", "CRUD for raw materials: name, quantity, unit, reorder_level, supplier"),
        ("FR-INV-02", "Flag items where quantity <= reorder_level as low stock"),
        ("FR-INV-03", "Inventory analytics: trend charts, low-stock summary"),
    ]),
    ("3.6 AI Copilot", [
        ("FR-AI-01", "WebSocket /ai/ws streams responses token-by-token in real time"),
        ("FR-AI-02", "Intent classification: GPT-4o (primary) → keyword rules → Ollama (fallback)"),
        ("FR-AI-03", "18+ intents: general_status, get_order_status, shipment_status, production_line_status, ask_manual, ask_policy, inventory_check, defect_analysis, count_orders, etc."),
        ("FR-AI-04", "Pre-routing override guards correct LLM misclassifications before dispatch"),
        ("FR-AI-05", "Voice input via Web Speech API; TTS output via browser SpeechSynthesis"),
        ("FR-AI-06", "Intent badge and confidence score shown in UI for transparency"),
    ]),
    ("3.7 RAG Knowledge Base", [
        ("FR-RAG-01", "Upload PDF, DOCX, TXT, XLSX documents up to 50MB"),
        ("FR-RAG-02", "Documents chunked (512 words, 64 overlap), embedded with text-embedding-3-small"),
        ("FR-RAG-03", "Query uses cosine similarity (pgvector); ILIKE keyword fallback if no embeddings"),
        ("FR-RAG-04", "Answers generated by GPT-4o → Ollama fallback → formatted excerpt fallback"),
        ("FR-RAG-05", "Responses include source citations: document name + page number"),
        ("FR-RAG-06", "Re-index button regenerates all chunks and embeddings for a document"),
    ]),
    ("3.8 Analytics", [
        ("FR-ANA-01", "Dashboard KPIs: active orders, shipments in transit, low stock, on-time delivery %"),
        ("FR-ANA-02", "Order trend chart (last 30 days), production efficiency per line"),
        ("FR-ANA-03", "AI generates and executes SQL for count/sum/trend questions"),
    ]),
]

for mod_title, reqs in modules:
    story.append(Paragraph(mod_title, S["H2"]))
    req_data = [[B("ID"), B("Requirement")]]
    for rid, rdesc in reqs:
        req_data.append([Paragraph(f"<font name='Courier'>{rid}</font>", S["body"]),
                         Paragraph(rdesc, S["body"])])
    story += [make_table(req_data, [3.2*cm, 14.3*cm]), Spacer(1, 0.2*cm)]

story.append(PageBreak())

# ── 4. Non-Functional Requirements ────────────────────────────
story += [
    Paragraph("4. Non-Functional Requirements", S["H1"]), HR(),
    Paragraph("4.1 Performance", S["H2"]),
    Paragraph("• REST API p95 response time < 500ms (non-AI endpoints)", S["bul"]),
    Paragraph("• Intent classification < 2s (GPT-4o) / < 8s (Ollama)", S["bul"]),
    Paragraph("• WebSocket first-token latency < 1.5s", S["bul"]),
    Paragraph("• Vector search < 200ms for 100K chunks with pgvector IVFFlat index", S["bul"]),
    Paragraph("4.2 Security", S["H2"]),
    Paragraph("• Passwords hashed with bcrypt (cost factor 12)", S["bul"]),
    Paragraph("• JWT access tokens expire 30 min; refresh tokens 7 days with rotation", S["bul"]),
    Paragraph("• All DB queries parameterised — no string concatenation (SQL injection prevention)", S["bul"]),
    Paragraph("• File uploads validated by MIME type and extension before processing", S["bul"]),
    Paragraph("4.3 Reliability", S["H2"]),
    Paragraph("• AI Copilot degrades gracefully: GPT-4o quota → keyword rules → Ollama", S["bul"]),
    Paragraph("• RAG falls back to ILIKE keyword search if embeddings unavailable", S["bul"]),
    Paragraph("• asyncpg connection pooling (pool_size=10)", S["bul"]),
    Paragraph("4.4 Scalability", S["H2"]),
    Paragraph("• Stateless FastAPI instances — horizontally scalable behind load balancer", S["bul"]),
    Paragraph("• pgvector IVFFlat index for sub-linear nearest-neighbour search", S["bul"]),
    Paragraph("4.5 Maintainability", S["H2"]),
    Paragraph("• Alembic migrations for all schema changes", S["bul"]),
    Paragraph("• Pydantic v2 schemas validate all request/response contracts", S["bul"]),
    Spacer(1, 0.3*cm), PageBreak(),
]

# ── 5. System Architecture ────────────────────────────────────
story += [
    Paragraph("5. System Architecture", S["H1"]), HR(),
    Paragraph("5.1 High-Level Diagram", S["H2"]),
    Paragraph("Browser (React/TS)  ←HTTP/WS→  FastAPI + Uvicorn  ←asyncpg→  PostgreSQL + pgvector", S["code"]),
    Paragraph("                                      ↕ HTTPS\n"
              "                            [OpenAI GPT-4o]   [Ollama Local LLM]", S["code"]),
    Paragraph("5.2 AI Intent Pipeline", S["H2"]),
    Paragraph("<b>Stage 1 — detect_intent():</b> GPT-4o (8s timeout) → keyword rules (instant) → Ollama", S["bul"]),
    Paragraph("<b>Stage 2 — route_intent():</b> 4 pre-routing override guards → domain handler", S["bul"]),
    Paragraph("<b>Stage 3 — Handler:</b> Service layer → DB query / RAG / SQL generation", S["bul"]),
    Paragraph("5.3 RAG Pipeline", S["H2"]),
    Paragraph("<b>Ingest:</b> upload → extract (PyMuPDF/docx/openpyxl) → chunk (512w/64 overlap) → embed → store in document_chunks", S["bul"]),
    Paragraph("<b>Query:</b> embed question → cosine similarity (MIN_SIM=0.72, TOP_K=5) → keyword fallback → GPT-4o answer → citations", S["bul"]),
    Spacer(1, 0.3*cm), PageBreak(),
]

# ── 6. Future Work ────────────────────────────────────────────
story += [Paragraph("6. Future Work & Recommendations", S["H1"]), HR()]

def future_section(title, items):
    out = [Paragraph(title, S["H2"])]
    data = [[B("Feature"), B("Description"), B("Priority")]]
    for feat, desc, pri in items:
        pri_color = "#DC2626" if pri == "High" else "#D97706" if pri == "Medium" else "#64748B"
        data.append([Paragraph(feat, S["body"]),
                     Paragraph(desc, S["body"]),
                     Paragraph(f'<font color="{pri_color}"><b>{pri}</b></font>', S["body"])])
    out.append(make_table(data, [4*cm, 11*cm, 2.5*cm], hdr_fill=TEAL))
    out.append(Spacer(1, 0.3*cm))
    return out

story += future_section("6.1 Short-Term (Next 3 Months)", [
    ("Bangla/Hindi UI",     "Multi-language support for local factory workers; voice in native language", "High"),
    ("Push Notifications",  "Real-time alerts for order delays, low stock, shipment updates", "High"),
    ("Mobile PWA",          "Responsive mobile layout + offline mode with service workers", "High"),
    ("Export Reports",      "One-click PDF/Excel export of production reports and shipment summaries", "Medium"),
    ("Bulk Order Import",   "CSV/Excel upload to create multiple production orders at once", "Medium"),
    ("Domain Fine-Tuning",  "Fine-tune GPT-4o on garment domain terms for higher intent accuracy", "Medium"),
])
story += future_section("6.2 Medium-Term (3–9 Months)", [
    ("Supplier Portal",     "External portal for suppliers to update material delivery status", "High"),
    ("Costing Module",      "Track material + labour costs, calculate margin per production order", "High"),
    ("Barcode / QR Scan",   "Mobile camera scan for order tracking and inventory check-in", "Medium"),
    ("Demand Forecasting",  "ML model to predict production demand and delivery delay probability", "Medium"),
    ("Email Alerts",        "Automated PO confirmations, dispatch alerts via SendGrid/SMTP", "Medium"),
    ("Multi-Currency",      "USD, BDT, EUR support for international buyer orders", "Low"),
])
story += future_section("6.3 Long-Term (9+ Months)", [
    ("ERP Integrations",    "API connectors to QuickBooks, SAP, Xero for accounting sync", "Medium"),
    ("CV Quality Control",  "AI defect detection from conveyor belt cameras using YOLO v8", "Medium"),
    ("Offline-First LLM",   "Fine-tuned Llama 3 on garment SOP corpus for fully offline AI", "Low"),
    ("Multi-Factory",       "Support factory groups with consolidated cross-location reporting", "Medium"),
    ("Audit Trail",         "Immutable change history for ISO 9001 / BSCI compliance", "High"),
    ("IoT Integration",     "Machine sensors for real-time efficiency and downtime data", "Low"),
])
story.append(PageBreak())

# ── 7. Appendix ───────────────────────────────────────────────
story += [Paragraph("7. Appendix — Full Tech Stack", S["H1"]), HR()]
stack_rows = [
    [B("Layer"), B("Technology"), B("Version"), B("Purpose")],
    ["Backend", "FastAPI",          "0.110+",   "REST API + WebSocket server"],
    ["Backend", "SQLAlchemy",       "2.0+",     "Async ORM + query builder"],
    ["Backend", "Alembic",          "1.13+",    "Database schema migrations"],
    ["Database","PostgreSQL",       "16+",      "Primary relational data store"],
    ["Database","pgvector",         "0.2.5+",   "1536-dim vector similarity search"],
    ["AI",      "OpenAI GPT-4o",    "Latest",   "Intent classification + RAG answer generation"],
    ["AI",      "Ollama llama3.1",  "3.1:8b",   "Local LLM fallback (no API key needed)"],
    ["AI",      "text-emb-3-small", "Latest",   "Document chunk embeddings (1536 dims)"],
    ["AI",      "nomic-embed-text", "Local",    "Offline embedding fallback"],
    ["Frontend","React",            "18",       "UI component framework"],
    ["Frontend","TypeScript",       "5+",       "Type-safe JavaScript"],
    ["Frontend","Vite",             "5+",       "Build tool + HMR dev server"],
    ["Frontend","Tailwind CSS",     "v4",       "Utility-first CSS framework"],
    ["Frontend","Zustand",          "4+",       "Lightweight global state management"],
    ["Auth",    "python-jose",      "3.3+",     "JWT encoding/decoding"],
    ["Auth",    "passlib + bcrypt", "1.7+",     "Secure password hashing"],
    ["DevOps",  "Docker Compose",   "Latest",   "Full-stack containerisation"],
    ["DevOps",  "Nginx",            "Latest",   "Production reverse proxy"],
]
story.append(make_table([[Paragraph(c, S["body"]) for c in row] for row in stack_rows],
                        [3*cm, 4*cm, 2.5*cm, 8*cm]))

# Build PDF
pdf_doc = SimpleDocTemplate(
    str(OUT / "Clavra_ProdPilot_SRS.pdf"), pagesize=A4,
    rightMargin=2*cm, leftMargin=2*cm, topMargin=2.5*cm, bottomMargin=2*cm
)
pdf_doc.build(story)
print("[OK]  Clavra_ProdPilot_SRS.pdf")

print("\nAll documents generated in:", OUT)
