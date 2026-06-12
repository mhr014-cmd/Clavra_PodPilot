"""Regenerate project_summary.docx and Clavra_ProdPilot_SRS.pdf with AI Vision section added."""
from pathlib import Path

OUT = Path("D:/projects/Claude/Clavra_ProdPilot/docs")
OUT.mkdir(exist_ok=True)

# ═══════════════════════════════════════════════════════════════
# DOCX HELPERS
# ═══════════════════════════════════════════════════════════════
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def shade_cell(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shd)

def hdr_row(table, headers, fill="1E40AF"):
    row = table.rows[0]
    for i, h in enumerate(headers):
        row.cells[i].text = h
        for run in row.cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
        shade_cell(row.cells[i], fill)

def add_row(table, *vals, key_bold=True, key_shade="DCE8F8"):
    row = table.add_row()
    for i, v in enumerate(vals):
        row.cells[i].text = str(v)
    if key_bold:
        for run in row.cells[0].paragraphs[0].runs:
            run.bold = True
    if key_shade:
        shade_cell(row.cells[0], key_shade)
    return row

def dset_heading(doc, text, level=1, rgb=(30, 64, 175)):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(*rgb)
    return h

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


# ═══════════════════════════════════════════════════════════════
# 1. PROJECT SUMMARY — DOCX (full regeneration + AI Vision)
# ═══════════════════════════════════════════════════════════════
doc = Document()

title = doc.add_heading("Clavra ProdPilot™", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in title.runs:
    r.font.color.rgb = RGBColor(30, 64, 175)
    r.font.size = Pt(28)

sub = doc.add_paragraph("AI-Powered Manufacturing ERP — Project Summary  |  v1.1  |  June 2026")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in sub.runs:
    r.font.color.rgb = RGBColor(100, 116, 139)
    r.italic = True
doc.add_paragraph()

# ── 1. Brief Summary ──────────────────────────────────────────
dset_heading(doc, "1. Brief Summary")
doc.add_paragraph(
    "Clavra ProdPilot™ is a full-stack, AI-driven ERP system purpose-built for garment manufacturing. "
    "It replaces WhatsApp and Excel-based workflows with a real-time dashboard, an intelligent AI Copilot "
    "(GPT-4o + Ollama), a RAG-powered Knowledge Base, and an AI Vision module for automated quality "
    "inspection — enabling factory managers to query production status, monitor shipments, manage inventory, "
    "retrieve SOPs, and detect fabric defects using natural language or voice."
)

# ── 2. Problem Statement ──────────────────────────────────────
dset_heading(doc, "2. Problem Statement")
doc.add_paragraph(
    "Small and mid-size garment factories in South & South-East Asia rely on WhatsApp messages and "
    "Excel sheets to manage production orders, shipments, and raw material inventory. Quality control "
    "is done manually with paper checklists, and defect images are shared over chat with no traceability. "
    "Clavra ProdPilot™ delivers an affordable, AI-first ERP where workers simply ask questions in plain "
    "language and photograph defects directly in the chat interface."
)

# ── 3. Key Features ───────────────────────────────────────────
dset_heading(doc, "3. Key Features")
features = [
    ("AI Copilot (18+ Intents)",
     "GPT-4o intent classification with keyword + Ollama fallbacks; voice input via Web Speech API"),
    ("AI Vision — Defect Analysis",
     "Upload or photograph garments/fabric; GPT-4o Vision (primary) or Ollama moondream/llava (local) "
     "detects defects, reads labels, inspects equipment. Returns severity, findings, and recommendations. "
     "Results stored in vision_analyses table with full audit trail."),
    ("Production Management",
     "Track orders, progress (produced vs target), defect qty, delivery dates with inline editing"),
    ("Shipment Tracking",
     "Real-time status lifecycle linked to production orders"),
    ("RAG Knowledge Base",
     "Upload PDFs/DOCX/XLSX SOPs; semantic vector search answers with citations"),
    ("Inventory Management",
     "Material stock, reorder alerts, supplier info"),
    ("Production Lines",
     "Line A/B/C/D assignment, status, and efficiency tracking"),
    ("Quality Control",
     "Defect analysis, QC records, inspection reports"),
    ("Analytics Dashboard",
     "KPI cards, order trend charts, inventory analytics"),
    ("Multi-tenant RBAC",
     "Org-scoped data with Admin/Manager/Operator roles"),
    ("WebSocket Streaming",
     "Real-time AI responses streamed token-by-token with intent badges"),
]
t = doc.add_table(rows=1, cols=2)
t.style = "Table Grid"
hdr_row(t, ["Feature", "Description"])
for name, desc in features:
    add_row(t, name, desc)

# ── 4. Technology Stack ───────────────────────────────────────
doc.add_paragraph()
dset_heading(doc, "4. Technology Stack")
stack = [
    ("Backend", "Python 3.11+, FastAPI, SQLAlchemy 2.0 (async), Alembic, Uvicorn"),
    ("Database", "PostgreSQL 16 + pgvector (1536-dim vector embeddings)"),
    ("Frontend", "React 18, TypeScript, Vite 5, Tailwind CSS v4, Zustand"),
    ("AI — Language", "OpenAI GPT-4o (intent + RAG answers), Ollama llama3.1:8b (local fallback)"),
    ("AI — Vision", "OpenAI GPT-4o Vision (primary), Ollama moondream:latest (local), Ollama llava (fallback)"),
    ("AI — Embeddings", "text-embedding-3-small (1536d, cloud), nomic-embed-text (local fallback)"),
    ("Auth", "JWT (python-jose), bcrypt, refresh token rotation"),
    ("Voice", "Web Speech API (browser-native STT + TTS)"),
    ("DevOps", "Docker Compose, Nginx reverse proxy, Alembic migrations"),
]
t2 = doc.add_table(rows=1, cols=2)
t2.style = "Table Grid"
hdr_row(t2, ["Layer", "Technologies"], fill="0F766E")
for layer, tech in stack:
    add_row(t2, layer, tech, key_shade="CCFBF1")

# ── 5. Architecture Overview ──────────────────────────────────
doc.add_paragraph()
dset_heading(doc, "5. Architecture Overview")
doc.add_paragraph(
    "Three-tier architecture:\n"
    "  Browser (React/TS) ←→ FastAPI REST + WebSocket ←→ PostgreSQL + pgvector\n\n"
    "AI intent pipeline (three stages):\n"
    "  Stage 1 — detect_intent():  GPT-4o (8s timeout) → keyword rules → Ollama fallback\n"
    "  Stage 2 — route_intent():   4 override guards correct misclassifications → dispatch\n"
    "  Stage 3 — Handler:          DB query / RAG / SQL generator / Vision service\n\n"
    "WebSocket sends intent badge immediately after Stage 1, then streams handler tokens."
)

# ── 6. AI Vision Module (NEW) ─────────────────────────────────
doc.add_paragraph()
dset_heading(doc, "6. AI Vision Module", rgb=(124, 45, 18))

intro = doc.add_paragraph(
    "The AI Vision module allows shop-floor workers to photograph garments, fabric rolls, "
    "labels, or production line equipment directly inside the AI Copilot chat. The image is "
    "analyzed automatically and a structured quality report is returned with findings, severity "
    "rating, and actionable recommendations."
)

doc.add_paragraph()
dset_heading(doc, "6.1  Supported Analysis Types", level=2, rgb=(124, 45, 18))
analysis_types = [
    ("defect_detection",
     "Identifies visible defects on fabric or finished garments: holes, stains, seam failures, "
     "colour inconsistencies, pilling, snags"),
    ("label_reading",
     "Reads and validates care labels, size labels, country-of-origin labels for compliance checks"),
    ("line_photo",
     "Assesses overall production line or workstation — worker ergonomics, machine setup, "
     "material placement"),
    ("equipment_check",
     "Inspects sewing machines, cutting equipment, or pressing irons for visible wear or "
     "maintenance needs"),
]
t3 = doc.add_table(rows=1, cols=2)
t3.style = "Table Grid"
hdr_row(t3, ["Analysis Type", "Description"], fill="7C2D12")
for atype, adesc in analysis_types:
    add_row(t3, atype, adesc, key_shade="FEF3C7")

doc.add_paragraph()
dset_heading(doc, "6.2  AI Model Fallback Chain", level=2, rgb=(124, 45, 18))
chain_data = [
    ("1st", "GPT-4o Vision (OpenAI)",
     "Returns structured JSON with findings, severity, defect_rate_estimate, confidence. "
     "Requires valid OpenAI API key with vision credits."),
    ("2nd", "Ollama moondream:latest",
     "Free local vision model (~1.6 GB). Runs offline. Returns parsed structured response. "
     "Pull with: ollama pull moondream"),
    ("3rd", "Ollama llava:latest",
     "Secondary local fallback if moondream not installed. Larger model (~4 GB). "
     "Pull with: ollama pull llava"),
    ("4th", "Graceful fallback",
     "Returns a user-friendly message with instructions to enable vision models. "
     "No crash — always returns a usable response."),
]
t4 = doc.add_table(rows=1, cols=3)
t4.style = "Table Grid"
hdr_row(t4, ["Priority", "Model", "Behaviour"], fill="7C2D12")
for pri, model, behav in chain_data:
    add_row(t4, pri, model, behav, key_shade="FEF3C7")

doc.add_paragraph()
dset_heading(doc, "6.3  Vision Response Structure", level=2, rgb=(124, 45, 18))
doc.add_paragraph(
    "Every vision analysis returns a standardised JSON object regardless of which model ran:"
)
p = doc.add_paragraph()
rn = p.add_run(
    '{\n'
    '  "analysis_type":       "defect_detection | label_reading | line_photo | equipment_check",\n'
    '  "findings":            ["finding 1", "finding 2", ...],\n'
    '  "severity":            "critical | major | minor | none",\n'
    '  "recommendations":     ["action 1", "action 2", ...],\n'
    '  "defect_rate_estimate": 0.05,\n'
    '  "confidence":          0.92,\n'
    '  "summary":             "Human-readable 2-sentence summary"\n'
    '}'
)
rn.font.name = "Courier New"
rn.font.size = Pt(9)

doc.add_paragraph()
dset_heading(doc, "6.4  Frontend Integration", level=2, rgb=(124, 45, 18))
doc.add_paragraph(
    "The AI Copilot chat includes an image upload button (camera icon) beside the message input. "
    "Workers tap the button, select or photograph an image, and the analysis result appears as a "
    "structured chat bubble alongside the conversation. The ImageUploadZone component accepts any "
    "image format (JPEG, PNG, WEBP) via the browser file picker."
)

doc.add_paragraph()
dset_heading(doc, "6.5  Data Persistence", level=2, rgb=(124, 45, 18))
db_cols = [
    ("id", "Primary key"),
    ("org_id", "Organisation scope (multi-tenant)"),
    ("user_id", "Analyst who uploaded the image"),
    ("image_filename", "Original uploaded filename"),
    ("analysis_type", "defect_detection / label_reading / line_photo / equipment_check"),
    ("findings", "JSON array of individual findings"),
    ("defect_rate", "Estimated defect rate (0.0 – 1.0)"),
    ("confidence", "Model confidence score (0.0 – 1.0)"),
    ("recommended_action", "Concatenated recommendation text"),
    ("summary", "Human-readable summary"),
    ("created_at", "Timestamp of analysis"),
]
t5 = doc.add_table(rows=1, cols=2)
t5.style = "Table Grid"
hdr_row(t5, ["Column (vision_analyses table)", "Description"], fill="7C2D12")
for col, desc in db_cols:
    add_row(t5, col, desc, key_shade="FEF3C7")

# ── 7. Development Access ─────────────────────────────────────
doc.add_paragraph()
dset_heading(doc, "7. Development Access")
for k, v in [("Frontend", "http://localhost:5174"),
             ("Backend API", "http://localhost:8000"),
             ("API Docs", "http://localhost:8000/docs"),
             ("Admin Login", "admin@clavra.com / Admin@123"),
             ("Ollama", "http://localhost:11434"),
             ("Vision models", "ollama pull moondream  (local vision, ~1.6 GB)")]:
    p = doc.add_paragraph()
    p.add_run(f"{k}: ").bold = True
    p.add_run(v)

doc.save(OUT / "project_summary.docx")
print("[OK]  project_summary.docx")


# ═══════════════════════════════════════════════════════════════
# 2. SRS PDF — full regeneration + FR-VIS section
# ═══════════════════════════════════════════════════════════════
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                TableStyle, HRFlowable, PageBreak)
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY

BLUE   = colors.HexColor("#1E40AF")
TEAL   = colors.HexColor("#0F766E")
ORANGE = colors.HexColor("#7C2D12")
AMBER  = colors.HexColor("#FEF3C7")
GRAY   = colors.HexColor("#64748B")
LGRAY  = colors.HexColor("#F1F5F9")
DGRAY  = colors.HexColor("#334155")
WHITE  = colors.white

styles = getSampleStyleSheet()

S = {
    "H0":   ParagraphStyle("H0",   parent=styles["Title"],   fontSize=26, textColor=BLUE, spaceAfter=6, alignment=TA_CENTER),
    "H0s":  ParagraphStyle("H0s",  parent=styles["Normal"],  fontSize=11, textColor=GRAY, spaceAfter=20, alignment=TA_CENTER),
    "H1":   ParagraphStyle("H1",   parent=styles["Heading1"],fontSize=15, textColor=BLUE, spaceBefore=18, spaceAfter=8),
    "H2":   ParagraphStyle("H2",   parent=styles["Heading2"],fontSize=12, textColor=TEAL, spaceBefore=12, spaceAfter=6),
    "H2v":  ParagraphStyle("H2v",  parent=styles["Heading2"],fontSize=12, textColor=ORANGE, spaceBefore=12, spaceAfter=6),
    "body": ParagraphStyle("body", parent=styles["Normal"],  fontSize=10, textColor=DGRAY, leading=15, spaceAfter=6, alignment=TA_JUSTIFY),
    "bul":  ParagraphStyle("bul",  parent=styles["Normal"],  fontSize=10, textColor=DGRAY, leading=14, spaceAfter=3, leftIndent=18),
    "code": ParagraphStyle("code", parent=styles["Normal"],  fontSize=8.5, fontName="Courier",
                           backColor=colors.HexColor("#F8FAFC"), leftIndent=14, rightIndent=14,
                           spaceAfter=6, leading=12, textColor=DGRAY),
}

def HR(): return HRFlowable(width="100%", thickness=1, color=colors.HexColor("#CBD5E1"), spaceAfter=6)
def B(t): return f"<b>{t}</b>"
def BO(t): return f'<font color="#7C2D12"><b>{t}</b></font>'

def make_table(data, col_widths, hdr_fill=BLUE):
    t = Table(data, colWidths=col_widths)
    t.setStyle(TableStyle([
        ("BACKGROUND",    (0,0), (-1,0),  hdr_fill),
        ("TEXTCOLOR",     (0,0), (-1,0),  WHITE),
        ("FONTNAME",      (0,0), (-1,0),  "Helvetica-Bold"),
        ("ROWBACKGROUNDS",(0,1), (-1,-1), [WHITE, LGRAY]),
        ("FONTSIZE",      (0,0), (-1,-1), 9.5),
        ("GRID",          (0,0), (-1,-1), 0.5, colors.HexColor("#CBD5E1")),
        ("VALIGN",        (0,0), (-1,-1), "TOP"),
        ("PADDING",       (0,0), (-1,-1), 5),
    ]))
    return t

def make_vision_table(data, col_widths):
    t = Table(data, colWidths=col_widths)
    t.setStyle(TableStyle([
        ("BACKGROUND",    (0,0), (-1,0),  ORANGE),
        ("TEXTCOLOR",     (0,0), (-1,0),  WHITE),
        ("FONTNAME",      (0,0), (-1,0),  "Helvetica-Bold"),
        ("ROWBACKGROUNDS",(0,1), (-1,-1), [WHITE, AMBER]),
        ("FONTSIZE",      (0,0), (-1,-1), 9.5),
        ("GRID",          (0,0), (-1,-1), 0.5, colors.HexColor("#FCD34D")),
        ("VALIGN",        (0,0), (-1,-1), "TOP"),
        ("PADDING",       (0,0), (-1,-1), 5),
    ]))
    return t

story = []

# ── Cover ──────────────────────────────────────────────────────
story += [
    Spacer(1, 1.5*cm),
    Paragraph("Clavra ProdPilot™", S["H0"]),
    Paragraph("Software Requirements Specification (SRS)", S["H0s"]),
    Paragraph("Version 1.1  |  June 2026  |  Includes AI Vision Module", S["H0s"]),
    HR(), Spacer(1, 0.5*cm),
]
doc_info = [
    [B("Document Title"), "Clavra ProdPilot™ — SRS"],
    [B("Version"), "1.1  (added: Section 3.9 AI Vision Analysis)"],
    [B("Date"), "June 12, 2026"],
    [B("Status"), "Active Development"],
    [B("Stack"), "FastAPI + React 18 + PostgreSQL + pgvector + GPT-4o + GPT-4o Vision + Ollama"],
    [B("Target Audience"), "Developers, QA Engineers, Stakeholders"],
]
story += [
    make_table([[Paragraph(a, S["body"]), Paragraph(b, S["body"])] for a, b in doc_info],
               [4.5*cm, 13*cm], hdr_fill=LGRAY),
    Spacer(1, 0.5*cm), PageBreak(),
]

# ── 1. Introduction ────────────────────────────────────────────
story += [
    Paragraph("1. Introduction", S["H1"]), HR(),
    Paragraph("1.1 Purpose", S["H2"]),
    Paragraph("This document specifies the functional and non-functional requirements for "
              "<b>Clavra ProdPilot™</b>, an AI-powered ERP for garment manufacturing. It is the "
              "authoritative reference for development, testing, and stakeholder alignment.", S["body"]),
    Paragraph("1.2 Scope", S["H2"]),
    Paragraph("The system covers: production order management, shipment tracking, raw material inventory, "
              "production line monitoring, quality control, AI Copilot (text + voice), RAG document retrieval, "
              "and <b>AI Vision</b> for automated quality inspection via image analysis.", S["body"]),
    Paragraph("1.3 Definitions", S["H2"]),
]
defs = [
    [B("Term"), B("Definition")],
    ["ERP",        "Enterprise Resource Planning — integrated management of core business processes"],
    ["RAG",        "Retrieval-Augmented Generation — LLM answers grounded in retrieved documents"],
    ["Intent",     "Classified purpose of a user message (e.g. shipment_status, defect_analysis)"],
    ["pgvector",   "PostgreSQL extension for storing and querying 1536-dim vector embeddings"],
    ["SOP",        "Standard Operating Procedure — step-by-step work instruction document"],
    ["RBAC",       "Role-Based Access Control — permissions tied to Admin/Manager/Operator roles"],
    ["JWT",        "JSON Web Token — stateless authentication token"],
    ["Ollama",     "Local LLM inference server (llama3.1:8b text, moondream/llava vision)"],
    ["moondream",  "Lightweight local vision model (~1.6 GB) for offline image analysis"],
    ["llava",      "LLaVA — secondary local vision model (~4 GB) used as fallback if moondream absent"],
    ["defect_analysis", "AI intent that routes to the Vision service for image-based QC inspection"],
]
story += [make_table([[Paragraph(a, S["body"]), Paragraph(b, S["body"])] for a, b in defs],
                     [3.5*cm, 14*cm]), Spacer(1, 0.3*cm), PageBreak()]

# ── 2. Overall Description ─────────────────────────────────────
story += [
    Paragraph("2. Overall Description", S["H1"]), HR(),
    Paragraph("2.1 Product Perspective", S["H2"]),
    Paragraph("Clavra ProdPilot™ is a standalone web application replacing spreadsheet-based factory "
              "workflows. It integrates with OpenAI for cloud AI (text + vision) and Ollama for local AI "
              "(text via llama3.1:8b, vision via moondream/llava), enabling fully offline deployments.", S["body"]),
    Paragraph("2.2 User Classes", S["H2"]),
]
users = [
    [B("User Class"), B("Role"), B("Primary Tasks")],
    ["Factory Admin",       "Full access", "Configure org, manage users, view all reports"],
    ["Production Manager",  "Manager",     "Monitor orders, update progress, view analytics"],
    ["QC Inspector",        "Operator",    "Upload defect images, log QC results, use AI Vision"],
    ["Line Supervisor",     "Operator",    "Update line status, record quantities"],
    ["Logistics Officer",   "Operator",    "Manage shipments, link orders"],
]
story += [make_table([[Paragraph(a, S["body"]), Paragraph(b, S["body"]), Paragraph(c, S["body"])]
                       for a, b, c in users],
                     [4.5*cm, 3.5*cm, 9.5*cm], hdr_fill=TEAL), Spacer(1, 0.3*cm)]
story += [
    Paragraph("2.3 Operating Environment", S["H2"]),
    Paragraph("• <b>Server:</b> Linux/Windows with Python 3.11+, PostgreSQL 16+, Ollama", S["bul"]),
    Paragraph("• <b>Client:</b> Chrome 120+ / Firefox 120+ / Edge 120+ with camera/microphone access", S["bul"]),
    Paragraph("• <b>Vision (cloud):</b> OpenAI GPT-4o Vision — requires API key with vision credits", S["bul"]),
    Paragraph("• <b>Vision (local):</b> Ollama moondream (~1.6 GB) or llava (~4 GB) — no API key needed", S["bul"]),
    Paragraph("• <b>Docker:</b> Full-stack via docker-compose + Nginx reverse proxy", S["bul"]),
    Spacer(1, 0.3*cm), PageBreak(),
]

# ── 3. Functional Requirements ─────────────────────────────────
story += [Paragraph("3. Functional Requirements", S["H1"]), HR()]

standard_modules = [
    ("3.1 Authentication & RBAC", [
        ("FR-AUTH-01", "Login with email+password returns JWT access token + refresh token"),
        ("FR-AUTH-02", "Refresh tokens rotate on each use; expired tokens rejected with HTTP 401"),
        ("FR-AUTH-03", "All endpoints (except /auth/login) require valid Bearer JWT"),
        ("FR-AUTH-04", "All data scoped to org_id — no cross-org data leakage"),
        ("FR-AUTH-05", "Roles: Admin > Manager > Operator — progressively fewer write permissions"),
    ]),
    ("3.2 Production Order Management", [
        ("FR-PROD-01", "CRUD for orders: order_no, buyer, style, colour, total_qty, status, delivery_date"),
        ("FR-PROD-02", "Status lifecycle: Draft → Cutting → Sewing → Finishing → QC → Packed → Shipped"),
        ("FR-PROD-03", "Record produced_qty and defect_qty; system calculates progress percentage"),
        ("FR-PROD-04", "Inline editing of progress, delivery date, and shipment link from UI"),
        ("FR-PROD-05", "Link order to shipment via PUT /orders/{id}/link-shipment"),
    ]),
    ("3.3 Production Line Management", [
        ("FR-LINE-01", "Define production lines (A/B/C/D) with capacity and current order assignment"),
        ("FR-LINE-02", "Real-time status: Idle / Running / Maintenance"),
        ("FR-LINE-03", "Track efficiency percentage per line"),
    ]),
    ("3.4 Shipment Management", [
        ("FR-SHIP-01", "CRUD for shipments: shipment_no, destination, carrier, status, estimated_delivery"),
        ("FR-SHIP-02", "Status lifecycle: Pending → In Transit → Delivered → Cancelled"),
        ("FR-SHIP-03", "Link/unlink production orders to shipments"),
    ]),
    ("3.5 Inventory Management", [
        ("FR-INV-01", "CRUD for raw materials: name, quantity, unit, reorder_level, supplier"),
        ("FR-INV-02", "Flag items where quantity <= reorder_level as low stock"),
        ("FR-INV-03", "Inventory analytics: trend charts, low-stock summary"),
    ]),
    ("3.6 AI Copilot", [
        ("FR-AI-01", "WebSocket /ai/ws streams responses token-by-token in real time"),
        ("FR-AI-02", "Intent pipeline: GPT-4o (primary) → keyword rules → Ollama (fallback)"),
        ("FR-AI-03", "18+ intents: general_status, get_order_status, shipment_status, "
                     "production_line_status, ask_manual, ask_policy, inventory_check, "
                     "defect_analysis, count_orders, and more"),
        ("FR-AI-04", "Pre-routing override guards correct LLM misclassifications before dispatch"),
        ("FR-AI-05", "Voice input via Web Speech API; TTS output via browser SpeechSynthesis"),
        ("FR-AI-06", "Intent badge and confidence score shown in UI for each response"),
    ]),
    ("3.7 RAG Knowledge Base", [
        ("FR-RAG-01", "Upload PDF, DOCX, TXT, XLSX documents up to 50 MB"),
        ("FR-RAG-02", "Documents chunked (512 words, 64 overlap), embedded with text-embedding-3-small"),
        ("FR-RAG-03", "Query: cosine similarity (pgvector, MIN_SIM=0.72); ILIKE fallback"),
        ("FR-RAG-04", "Answers: GPT-4o → Ollama fallback → formatted excerpt fallback"),
        ("FR-RAG-05", "Responses include source citations: document name + page number"),
        ("FR-RAG-06", "Re-index button regenerates all chunks and embeddings for a document"),
    ]),
    ("3.8 Analytics", [
        ("FR-ANA-01", "Dashboard KPIs: active orders, in-transit shipments, low stock, on-time %"),
        ("FR-ANA-02", "Order trend chart (last 30 days), production efficiency per line"),
        ("FR-ANA-03", "AI generates and executes SQL for count/sum/trend questions"),
    ]),
]

for mod_title, reqs in standard_modules:
    story.append(Paragraph(mod_title, S["H2"]))
    req_data = [[B("ID"), B("Requirement")]]
    for rid, rdesc in reqs:
        req_data.append([Paragraph(f"<font name='Courier'>{rid}</font>", S["body"]),
                         Paragraph(rdesc, S["body"])])
    story += [make_table(req_data, [3.2*cm, 14.3*cm]), Spacer(1, 0.2*cm)]

# ── 3.9 AI VISION (highlighted in orange) ─────────────────────
story += [
    Paragraph("3.9 AI Vision Analysis", S["H2v"]),
    Paragraph(
        "The AI Vision module allows QC inspectors and shop-floor workers to photograph or upload "
        "images of garments, fabric, labels, or equipment directly inside the AI Copilot chat. "
        "Images are analyzed by a vision AI and a structured quality report is returned.", S["body"]),
]

vis_reqs = [
    [BO("ID"), BO("Requirement")],
    [Paragraph("<font name='Courier'>FR-VIS-01</font>", S["body"]),
     Paragraph("Users upload images (JPEG, PNG, WEBP) via the camera button in the AI Copilot chat interface", S["body"])],
    [Paragraph("<font name='Courier'>FR-VIS-02</font>", S["body"]),
     Paragraph("Image analysis triggered by <b>defect_analysis</b> intent — routed to vision_service.analyze_image()", S["body"])],
    [Paragraph("<font name='Courier'>FR-VIS-03</font>", S["body"]),
     Paragraph("Analysis pipeline: GPT-4o Vision (cloud, high detail) → Ollama moondream (~1.6 GB local) → Ollama llava (~4 GB local) → graceful fallback message", S["body"])],
    [Paragraph("<font name='Courier'>FR-VIS-04</font>", S["body"]),
     Paragraph("Supported analysis types: defect_detection | label_reading | line_photo | equipment_check", S["body"])],
    [Paragraph("<font name='Courier'>FR-VIS-05</font>", S["body"]),
     Paragraph("Every analysis returns structured JSON: analysis_type, findings[], severity (critical/major/minor/none), recommendations[], defect_rate_estimate, confidence, summary", S["body"])],
    [Paragraph("<font name='Courier'>FR-VIS-06</font>", S["body"]),
     Paragraph("GPT-4o Vision uses json_object response format with VISION_PROMPT for structured output; Ollama response parsed via _parse_ollama_response()", S["body"])],
    [Paragraph("<font name='Courier'>FR-VIS-07</font>", S["body"]),
     Paragraph("All analysis results persisted in vision_analyses table (org_id, user_id, image_filename, analysis_type, findings JSON, defect_rate, confidence, summary, created_at)", S["body"])],
    [Paragraph("<font name='Courier'>FR-VIS-08</font>", S["body"]),
     Paragraph("System discovers available Ollama vision models dynamically via GET /api/tags — tries moondream first, llava second", S["body"])],
    [Paragraph("<font name='Courier'>FR-VIS-09</font>", S["body"]),
     Paragraph("Graceful fallback: if no vision model available, returns user-friendly message with instructions to add OpenAI credits or pull moondream — never crashes", S["body"])],
    [Paragraph("<font name='Courier'>FR-VIS-10</font>", S["body"]),
     Paragraph("Ollama vision calls wrapped with asyncio.wait_for(timeout=60s) to prevent blocking the event loop on slow local inference", S["body"])],
]

story += [make_vision_table(vis_reqs, [3.2*cm, 14.3*cm]), Spacer(1, 0.3*cm)]

# Vision analysis types sub-table
story += [Paragraph("Supported Analysis Types Detail:", S["H2v"])]
vis_types = [
    [BO("Analysis Type"), BO("Trigger Keywords"), BO("What AI Looks For")],
    [Paragraph("defect_detection", S["body"]),
     Paragraph("defect, fault, damage, hole, stain, tear, seam, colour", S["body"]),
     Paragraph("Holes, stains, seam failures, colour inconsistencies, pilling, snags, weave defects", S["body"])],
    [Paragraph("label_reading", S["body"]),
     Paragraph("label, tag, care label, size, country of origin", S["body"]),
     Paragraph("Care instructions, size markings, country-of-origin, brand labels, compliance marks", S["body"])],
    [Paragraph("line_photo", S["body"]),
     Paragraph("production line, workstation, floor, line photo", S["body"]),
     Paragraph("Worker ergonomics, machine setup, material placement, workstation cleanliness", S["body"])],
    [Paragraph("equipment_check", S["body"]),
     Paragraph("machine, equipment, sewing machine, needle, iron, cutter", S["body"]),
     Paragraph("Visible wear, needle condition, calibration issues, maintenance needs", S["body"])],
]
story += [make_vision_table(vis_types, [3.5*cm, 5*cm, 9*cm]), Spacer(1, 0.3*cm)]
story.append(PageBreak())

# ── 4. Non-Functional Requirements ────────────────────────────
story += [
    Paragraph("4. Non-Functional Requirements", S["H1"]), HR(),
    Paragraph("4.1 Performance", S["H2"]),
    Paragraph("• REST API p95 response time < 500ms (non-AI endpoints)", S["bul"]),
    Paragraph("• Intent classification < 2s (GPT-4o) / < 8s (Ollama text)", S["bul"]),
    Paragraph("• WebSocket first-token latency < 1.5s", S["bul"]),
    Paragraph("• Vision analysis: GPT-4o Vision < 10s; Ollama moondream < 60s (local hardware dependent)", S["bul"]),
    Paragraph("• Vector similarity search < 200ms for 100K chunks", S["bul"]),
    Paragraph("4.2 Security", S["H2"]),
    Paragraph("• Passwords hashed with bcrypt (cost factor 12)", S["bul"]),
    Paragraph("• JWT access tokens expire 30 min; refresh tokens 7 days with rotation", S["bul"]),
    Paragraph("• All DB queries parameterised — SQL injection prevention", S["bul"]),
    Paragraph("• Uploaded images validated by MIME type before processing", S["bul"]),
    Paragraph("• Vision analysis results scoped to org_id — no cross-org leakage", S["bul"]),
    Paragraph("4.3 Reliability", S["H2"]),
    Paragraph("• AI Copilot: GPT-4o quota → keyword rules → Ollama text", S["bul"]),
    Paragraph("• AI Vision: GPT-4o Vision → Ollama moondream → Ollama llava → graceful message", S["bul"]),
    Paragraph("• RAG: vector search → ILIKE keyword fallback", S["bul"]),
    Paragraph("4.4 Scalability", S["H2"]),
    Paragraph("• Stateless FastAPI instances — horizontally scalable behind load balancer", S["bul"]),
    Paragraph("• pgvector IVFFlat index for sub-linear nearest-neighbour search", S["bul"]),
    Spacer(1, 0.3*cm), PageBreak(),
]

# ── 5. System Architecture ─────────────────────────────────────
story += [
    Paragraph("5. System Architecture", S["H1"]), HR(),
    Paragraph("5.1 High-Level Diagram", S["H2"]),
    Paragraph("Browser (React/TS)  ←HTTP/WS→  FastAPI + Uvicorn  ←asyncpg→  PostgreSQL + pgvector", S["code"]),
    Paragraph("                                        |HTTPS\n"
              "               [OpenAI GPT-4o Text]   [OpenAI GPT-4o Vision]\n"
              "               [Ollama llama3.1:8b]   [Ollama moondream / llava]", S["code"]),
    Paragraph("5.2 AI Text Intent Pipeline", S["H2"]),
    Paragraph("<b>Stage 1</b> — detect_intent(): GPT-4o (8s timeout) → keyword rules → Ollama", S["bul"]),
    Paragraph("<b>Stage 2</b> — route_intent(): 4 override guards → domain handler dispatch", S["bul"]),
    Paragraph("<b>Stage 3</b> — Handler: DB query / RAG / SQL generation / Vision service", S["bul"]),
    Paragraph("5.3 AI Vision Pipeline", S["H2v"]),
    Paragraph("<b>Step 1</b> — User uploads image via ImageUploadZone (camera button in chat)", S["bul"]),
    Paragraph("<b>Step 2</b> — defect_analysis intent routes to vision_service.analyze_image()", S["bul"]),
    Paragraph("<b>Step 3</b> — GPT-4o Vision tried first (json_object mode, high detail)", S["bul"]),
    Paragraph("<b>Step 4</b> — If OpenAI unavailable: _find_ollama_vision_model() queries /api/tags", S["bul"]),
    Paragraph("<b>Step 5</b> — moondream tried first, then llava, with 60s asyncio timeout", S["bul"]),
    Paragraph("<b>Step 6</b> — _parse_ollama_response() normalises free-text into structured dict", S["bul"]),
    Paragraph("<b>Step 7</b> — Result saved to vision_analyses table + returned as chat bubble", S["bul"]),
    Paragraph("5.4 RAG Pipeline", S["H2"]),
    Paragraph("<b>Ingest:</b> upload → extract (PyMuPDF/docx/openpyxl) → chunk (512w/64 overlap) → embed → pgvector", S["bul"]),
    Paragraph("<b>Query:</b> embed question → cosine similarity (MIN_SIM=0.72, TOP_K=5) → keyword fallback → answer", S["bul"]),
    Spacer(1, 0.3*cm), PageBreak(),
]

# ── 6. Future Work ─────────────────────────────────────────────
story += [Paragraph("6. Future Work & Recommendations", S["H1"]), HR()]

def future_section(title, items, style="H2"):
    s = S[style]
    out = [Paragraph(title, s)]
    data = [[B("Feature"), B("Description"), B("Priority")]]
    for feat, desc, pri in items:
        color = "#DC2626" if pri=="High" else "#D97706" if pri=="Medium" else "#64748B"
        data.append([Paragraph(feat, S["body"]), Paragraph(desc, S["body"]),
                     Paragraph(f'<font color="{color}"><b>{pri}</b></font>', S["body"])])
    out.append(make_table(data, [4*cm, 11*cm, 2.5*cm], hdr_fill=TEAL))
    out.append(Spacer(1, 0.3*cm))
    return out

story += future_section("6.1 Short-Term (Next 3 Months)", [
    ("Bangla/Hindi UI",      "Multi-language support for local factory workers", "High"),
    ("Push Notifications",   "Alerts for order delays, low stock, delivery updates", "High"),
    ("Mobile PWA",           "Responsive mobile + offline mode with service workers", "High"),
    ("Export Reports",       "PDF/Excel export of production reports and shipment summaries", "Medium"),
    ("Vision Batch Mode",    "Upload multiple defect images at once for bulk QC inspection", "Medium"),
    ("Vision History Page",  "Dedicated page showing all past vision analyses with filters", "Medium"),
])
story += future_section("6.2 Medium-Term (3–9 Months)", [
    ("Supplier Portal",      "External portal for suppliers to update material delivery status", "High"),
    ("Costing Module",       "Material + labour cost tracking, margin per production order", "High"),
    ("Barcode/QR Scan",      "Mobile camera for order tracking and inventory check-in", "Medium"),
    ("Vision Defect Stats",  "Dashboard chart: defect rates by order, line, garment type over time", "Medium"),
    ("Conveyor Camera",      "RTSP camera stream integration for automated line QC (no manual upload)", "Medium"),
    ("Email Alerts",         "Automated PO confirmations, dispatch alerts via SendGrid", "Medium"),
])
story += future_section("6.3 Long-Term (9+ Months)", [
    ("YOLO v8 Integration",  "Fine-tuned object detection for specific defect types on conveyor belt", "Medium"),
    ("ERP Integrations",     "QuickBooks, SAP, Xero connectors for accounting sync", "Medium"),
    ("Audit Trail",          "Immutable change history for ISO 9001 / BSCI compliance", "High"),
    ("IoT + Vision",         "Machine sensor data combined with camera feeds for real-time QC", "Low"),
    ("Vision Model Training","Fine-tune moondream/llava on garment defect dataset for higher accuracy", "Low"),
])
story.append(PageBreak())

# ── 7. Appendix ────────────────────────────────────────────────
story += [Paragraph("7. Appendix — Full Tech Stack", S["H1"]), HR()]
stack_rows = [
    [B("Layer"), B("Technology"), B("Version"), B("Purpose")],
    ["Backend",  "FastAPI",             "0.110+",  "REST API + WebSocket server"],
    ["Backend",  "SQLAlchemy",          "2.0+",    "Async ORM + query builder"],
    ["Backend",  "Alembic",             "1.13+",   "Database schema migrations"],
    ["Database", "PostgreSQL",          "16+",     "Primary relational data store"],
    ["Database", "pgvector",            "0.2.5+",  "1536-dim vector similarity search"],
    ["AI Text",  "OpenAI GPT-4o",       "Latest",  "Intent classification + RAG answer generation"],
    ["AI Vision","OpenAI GPT-4o Vision","Latest",  "Image defect detection, label reading, QC inspection"],
    ["AI Vision","Ollama moondream",    "latest",  "Local vision fallback (~1.6 GB, offline capable)"],
    ["AI Vision","Ollama llava",        "latest",  "Secondary local vision fallback (~4 GB)"],
    ["AI Text",  "Ollama llama3.1",     "8b",      "Local text LLM fallback"],
    ["AI Embed", "text-emb-3-small",    "Latest",  "Document chunk embeddings (1536 dims)"],
    ["AI Embed", "nomic-embed-text",    "Local",   "Offline embedding fallback"],
    ["Frontend", "React",               "18",      "UI component framework"],
    ["Frontend", "TypeScript",          "5+",      "Type-safe JavaScript"],
    ["Frontend", "Vite",                "5+",      "Build tool + HMR dev server"],
    ["Frontend", "Tailwind CSS",        "v4",      "Utility-first CSS framework"],
    ["Frontend", "Zustand",             "4+",      "Global state management"],
    ["Auth",     "python-jose",         "3.3+",    "JWT encoding/decoding"],
    ["Auth",     "passlib + bcrypt",    "1.7+",    "Secure password hashing"],
    ["DevOps",   "Docker Compose",      "Latest",  "Full-stack containerisation"],
    ["DevOps",   "Nginx",               "Latest",  "Production reverse proxy"],
]
story.append(make_table([[Paragraph(c, S["body"]) for c in row] for row in stack_rows],
                        [2.5*cm, 4*cm, 2.5*cm, 8.5*cm]))

# Build PDF
pdf_doc = SimpleDocTemplate(
    str(OUT / "Clavra_ProdPilot_SRS.pdf"), pagesize=A4,
    rightMargin=2*cm, leftMargin=2*cm, topMargin=2.5*cm, bottomMargin=2*cm
)
pdf_doc.build(story)
print("[OK]  Clavra_ProdPilot_SRS.pdf")

print("\nDone. Files saved to:", OUT)
